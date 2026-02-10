#!/usr/bin/env python3
"""
Generate a professionally formatted Word document proposal for BCC Website EOI.
Job Angula Technology Consulting
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Colour Palette (Ocean / BCC-inspired) ──
DEEP_BLUE = RGBColor(0x00, 0x3D, 0x6B)      # Primary headings
OCEAN_BLUE = RGBColor(0x00, 0x7B, 0xB8)      # Secondary headings
LIGHT_BLUE = RGBColor(0xD6, 0xEA, 0xF8)      # Table header bg
ACCENT_GREEN = RGBColor(0x00, 0x9B, 0x72)     # Accents / highlights
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)        # Body text
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)      # Secondary text
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HEADER_BG = "0073B7"
TABLE_ALT_ROW = "EBF5FB"
ACCENT_BAR = "00639B"
HIGHLIGHT_BG = "E8F8F5"
SECTION_BG = "F0F7FF"

doc = Document()

# ── Page Margins ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Default Font ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = DARK_GRAY

# ── Helper Functions ──

def set_cell_shading(cell, color_hex):
    """Set background colour of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """Set borders on a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Style header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, TABLE_HEADER_BG)

    # Style data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK_GRAY
            if r_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_ROW)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(width)

    return table

def add_info_box(doc, title, content, bg_color=SECTION_BG):
    """Add a highlighted information box using a single-cell table."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, bg_color.replace("#", ""))

    borders = {"top": {"val": "single", "sz": "12", "color": ACCENT_BAR},
               "bottom": {"val": "single", "sz": "4", "color": ACCENT_BAR},
               "left": {"val": "single", "sz": "4", "color": ACCENT_BAR},
               "right": {"val": "single", "sz": "4", "color": ACCENT_BAR}}
    set_cell_border(cell, **borders)

    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DEEP_BLUE
    run.font.name = 'Calibri'

    p2 = cell.add_paragraph()
    run2 = p2.add_run(content)
    run2.font.size = Pt(10)
    run2.font.color.rgb = DARK_GRAY
    run2.font.name = 'Calibri'

    return table

def add_section_heading(doc, text, level=1):
    """Add a coloured section heading."""
    p = doc.add_paragraph()
    if level == 1:
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = DEEP_BLUE
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Add a bottom border line
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="{ACCENT_BAR}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
    elif level == 2:
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = OCEAN_BLUE
        run.font.name = 'Calibri'
    elif level == 3:
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = ACCENT_GREEN
        run.font.name = 'Calibri'
    return p

def add_body(doc, text):
    """Add body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GRAY
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(doc, text, bold_prefix=None):
    """Add a bullet point, optionally with bold prefix."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_GRAY
        run.font.name = 'Calibri'
        run2 = p.add_run(text)
        run2.font.size = Pt(11)
        run2.font.color.rgb = DARK_GRAY
        run2.font.name = 'Calibri'
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_GRAY
        run.font.name = 'Calibri'
    return p

def add_progress_bar(doc, label, percentage, color=ACCENT_BAR):
    """Add a visual progress/percentage bar using a table."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(4)
    table.columns[1].width = Cm(12)

    # Label cell
    cell0 = table.rows[0].cells[0]
    cell0.text = ""
    p0 = cell0.paragraphs[0]
    run0 = p0.add_run(label)
    run0.font.size = Pt(9)
    run0.font.color.rgb = DARK_GRAY
    run0.font.name = 'Calibri'
    run0.bold = True

    # Bar cell - use nested table to simulate progress bar
    cell1 = table.rows[0].cells[1]
    cell1.text = ""
    inner = cell1.add_table(rows=1, cols=2)
    filled_width = max(1, int(percentage / 100 * 12))
    empty_width = 12 - filled_width

    filled_cell = inner.rows[0].cells[0]
    filled_cell.width = Cm(filled_width)
    set_cell_shading(filled_cell, color)
    fp = filled_cell.paragraphs[0]
    fr = fp.add_run(f" {percentage}%")
    fr.font.size = Pt(8)
    fr.font.color.rgb = WHITE
    fr.font.name = 'Calibri'
    fr.bold = True

    empty_cell = inner.rows[0].cells[1]
    empty_cell.width = Cm(empty_width)
    set_cell_shading(empty_cell, "E8E8E8")
    empty_cell.text = ""

    return table

def add_icon_card_row(doc, cards):
    """Add a row of icon cards (emoji + title + description) using a table."""
    table = doc.add_table(rows=2, cols=len(cards))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (icon, title, desc) in enumerate(cards):
        # Icon + Title row
        cell_top = table.rows[0].cells[i]
        set_cell_shading(cell_top, SECTION_BG.replace("#", ""))
        cell_top.text = ""
        p1 = cell_top.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run_icon = p1.add_run(icon + " ")
        run_icon.font.size = Pt(18)

        run_title = p1.add_run(title)
        run_title.bold = True
        run_title.font.size = Pt(11)
        run_title.font.color.rgb = DEEP_BLUE
        run_title.font.name = 'Calibri'

        # Description row
        cell_bot = table.rows[1].cells[i]
        set_cell_shading(cell_bot, SECTION_BG.replace("#", ""))
        cell_bot.text = ""
        p2 = cell_bot.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_desc = p2.add_run(desc)
        run_desc.font.size = Pt(9)
        run_desc.font.color.rgb = MEDIUM_GRAY
        run_desc.font.name = 'Calibri'

    return table

def add_phase_block(doc, phase_num, title, duration, items, color=ACCENT_BAR):
    """Add a visually distinct project phase block."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(2)
    table.columns[1].width = Cm(14)

    # Phase number circle (simulated with colored cell)
    cell0 = table.rows[0].cells[0]
    set_cell_shading(cell0, color)
    cell0.text = ""
    p0 = cell0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run0 = p0.add_run(f"\n{phase_num}\n")
    run0.bold = True
    run0.font.size = Pt(20)
    run0.font.color.rgb = WHITE
    run0.font.name = 'Calibri'

    # Phase content
    cell1 = table.rows[0].cells[1]
    cell1.text = ""
    p1 = cell1.paragraphs[0]
    run_title = p1.add_run(title)
    run_title.bold = True
    run_title.font.size = Pt(12)
    run_title.font.color.rgb = DEEP_BLUE
    run_title.font.name = 'Calibri'

    p_dur = cell1.add_paragraph()
    run_dur = p_dur.add_run(duration)
    run_dur.font.size = Pt(9)
    run_dur.font.color.rgb = OCEAN_BLUE
    run_dur.font.name = 'Calibri'
    run_dur.italic = True

    for item in items:
        p_item = cell1.add_paragraph()
        run_bullet = p_item.add_run("  \u25B8  " + item)
        run_bullet.font.size = Pt(10)
        run_bullet.font.color.rgb = DARK_GRAY
        run_bullet.font.name = 'Calibri'

    return table


# ═══════════════════════════════════════════════════════════════════
#                        COVER PAGE
# ═══════════════════════════════════════════════════════════════════

# Top colour bar
p_bar = doc.add_paragraph()
pPr = p_bar._p.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:top w:val="single" w:sz="48" w:space="1" w:color="{ACCENT_BAR}"/>'
    f'</w:pBdr>'
)
pPr.append(pBdr)

# Spacer
doc.add_paragraph("")
doc.add_paragraph("")

# Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_title.add_run("EXPRESSION OF INTEREST")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = DEEP_BLUE
run.font.name = 'Calibri'

# Subtitle
p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p_sub.add_run("Revamping and Re-Designing of the\nBenguela Current Convention (BCC) Website")
run2.font.size = Pt(16)
run2.font.color.rgb = OCEAN_BLUE
run2.font.name = 'Calibri'

# Decorative line
p_line = doc.add_paragraph()
p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr2 = p_line._p.get_or_add_pPr()
pBdr2 = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:bottom w:val="single" w:sz="24" w:space="1" w:color="{ACCENT_BAR}"/>'
    f'</w:pBdr>'
)
pPr2.append(pBdr2)

doc.add_paragraph("")

# Submitted by box
cover_table = doc.add_table(rows=1, cols=2)
cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER

cell_left = cover_table.rows[0].cells[0]
cell_left.text = ""
p_sb = cell_left.paragraphs[0]
p_sb.alignment = WD_ALIGN_PARAGRAPH.LEFT
run_sb_label = p_sb.add_run("Submitted by:")
run_sb_label.font.size = Pt(10)
run_sb_label.font.color.rgb = MEDIUM_GRAY
run_sb_label.font.name = 'Calibri'

p_sb2 = cell_left.add_paragraph()
run_sb2 = p_sb2.add_run("Job Angula Technology Consulting")
run_sb2.bold = True
run_sb2.font.size = Pt(14)
run_sb2.font.color.rgb = DEEP_BLUE
run_sb2.font.name = 'Calibri'

p_sb3 = cell_left.add_paragraph()
run_sb3 = p_sb3.add_run("Windhoek, Namibia\njob@angulaconsulting.com\nwww.angulaconsulting.com")
run_sb3.font.size = Pt(10)
run_sb3.font.color.rgb = DARK_GRAY
run_sb3.font.name = 'Calibri'

cell_right = cover_table.rows[0].cells[1]
cell_right.text = ""
p_to = cell_right.paragraphs[0]
p_to.alignment = WD_ALIGN_PARAGRAPH.LEFT
run_to_label = p_to.add_run("Submitted to:")
run_to_label.font.size = Pt(10)
run_to_label.font.color.rgb = MEDIUM_GRAY
run_to_label.font.name = 'Calibri'

p_to2 = cell_right.add_paragraph()
run_to2 = p_to2.add_run("Benguela Current Convention (BCC)")
run_to2.bold = True
run_to2.font.size = Pt(14)
run_to2.font.color.rgb = DEEP_BLUE
run_to2.font.name = 'Calibri'

p_to3 = cell_right.add_paragraph()
run_to3 = p_to3.add_run("Secretariat\nNo. 1 Strand Street\nP/Bag 5031 Swakopmund, Namibia")
run_to3.font.size = Pt(10)
run_to3.font.color.rgb = DARK_GRAY
run_to3.font.name = 'Calibri'

doc.add_paragraph("")

# Date
p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_d = p_date.add_run("February 2026")
run_d.bold = True
run_d.font.size = Pt(14)
run_d.font.color.rgb = OCEAN_BLUE
run_d.font.name = 'Calibri'

# Bottom colour bar
p_bar2 = doc.add_paragraph()
pPr3 = p_bar2._p.get_or_add_pPr()
pBdr3 = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:bottom w:val="single" w:sz="48" w:space="1" w:color="{ACCENT_BAR}"/>'
    f'</w:pBdr>'
)
pPr3.append(pBdr3)

# ── Page Break ──
doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
#                    TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════

add_section_heading(doc, "Table of Contents")
doc.add_paragraph("")

toc_items = [
    ("Section A", "Cover Letter", "3"),
    ("Section B", "Company Profile", "4"),
    ("Section C", "Technical Proposal", "6"),
    ("Section D", "Work Plan", "14"),
    ("Section E", "Financial Proposal", "16"),
    ("Section F", "References", "18"),
    ("Annexures", "Supporting Documents", "19"),
]

toc_table = doc.add_table(rows=len(toc_items), cols=3)
toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (sec, title, pg) in enumerate(toc_items):
    cell0 = toc_table.rows[i].cells[0]
    cell0.text = ""
    r0 = cell0.paragraphs[0].add_run(sec)
    r0.bold = True
    r0.font.size = Pt(11)
    r0.font.color.rgb = OCEAN_BLUE
    r0.font.name = 'Calibri'

    cell1 = toc_table.rows[i].cells[1]
    cell1.text = ""
    r1 = cell1.paragraphs[0].add_run(title)
    r1.font.size = Pt(11)
    r1.font.color.rgb = DARK_GRAY
    r1.font.name = 'Calibri'

    cell2 = toc_table.rows[i].cells[2]
    cell2.text = ""
    cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = cell2.paragraphs[0].add_run(pg)
    r2.font.size = Pt(11)
    r2.font.color.rgb = MEDIUM_GRAY
    r2.font.name = 'Calibri'

    if i % 2 == 0:
        set_cell_shading(cell0, TABLE_ALT_ROW)
        set_cell_shading(cell1, TABLE_ALT_ROW)
        set_cell_shading(cell2, TABLE_ALT_ROW)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
#                    SECTION A: COVER LETTER
# ═══════════════════════════════════════════════════════════════════

add_section_heading(doc, "Section A: Cover Letter")
doc.add_paragraph("")

add_body(doc, "[DD] February 2026")
doc.add_paragraph("")

add_body(doc, "The Executive Secretary\nBenguela Current Convention (BCC) Secretariat\nNo. 1 Strand Street\nP/Bag 5031\nSwakopmund, Namibia")
doc.add_paragraph("")

p_re = doc.add_paragraph()
run_re = p_re.add_run("Re: Expression of Interest \u2013 Revamping and Re-Designing of the BCC Website")
run_re.bold = True
run_re.font.size = Pt(11)
run_re.font.color.rgb = DEEP_BLUE
run_re.font.name = 'Calibri'
doc.add_paragraph("")

add_body(doc, "Dear Sir/Madam,")
doc.add_paragraph("")

add_body(doc, "I, Job Angula, the Principal Consultant and sole proprietor of Job Angula Technology Consulting, hereby submit this Expression of Interest in response to the call for the design, development, implementation, and hosting of the Benguela Current Convention (BCC) website.")

add_body(doc, "Job Angula Technology Consulting is a Namibian-based technology consulting firm specialising in custom web application development, website design, content management systems, and digital solutions for organisations across Southern Africa. We have a demonstrated track record of delivering modern, responsive, and user-friendly web platforms for both private and public sector clients.")

add_body(doc, "We understand that the BCC requires a website that is interactive, modern, and reflective of its identity as the world\u2019s first intergovernmental Convention based on a multi-sectoral approach to Large Marine Ecosystem ocean governance. We are confident that our expertise in building custom, tailor-made content management systems and modern web applications positions us uniquely to deliver a solution that meets and exceeds the BCC\u2019s expectations.")

add_body(doc, "Our approach centres on building a bespoke Content Management System specifically engineered for BCC\u2019s operational needs \u2013 enabling your Secretariat staff to independently manage bilingual content (English and Portuguese), upload documents, manage media, and maintain the website without requiring technical expertise. This custom approach eliminates the security vulnerabilities and bloat associated with off-the-shelf CMS platforms while providing exactly the features BCC needs.")

add_body(doc, "We are committed to delivering a high-quality, secure, and fully functional website within the stipulated timelines, and to providing reliable hosting and maintenance services for the three-year contract period.")

add_body(doc, "We look forward to the opportunity to contribute to the BCC\u2019s digital transformation and enhanced stakeholder engagement.")
doc.add_paragraph("")

add_body(doc, "Yours faithfully,")
doc.add_paragraph("")
doc.add_paragraph("")

p_sig = doc.add_paragraph()
run_sig = p_sig.add_run("___________________________")
run_sig.font.color.rgb = MEDIUM_GRAY
p_sig2 = doc.add_paragraph()
run_name = p_sig2.add_run("Job Angula")
run_name.bold = True
run_name.font.size = Pt(12)
run_name.font.color.rgb = DEEP_BLUE
run_name.font.name = 'Calibri'

add_body(doc, "Principal Consultant\nJob Angula Technology Consulting\nWindhoek, Namibia\nPhone: +264 XX XXX XXXX\nEmail: job@angulaconsulting.com")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
#                    SECTION B: COMPANY PROFILE
# ═══════════════════════════════════════════════════════════════════

add_section_heading(doc, "Section B: Company Profile")
doc.add_paragraph("")

add_section_heading(doc, "1. About the Firm", level=2)
add_body(doc, "Job Angula Technology Consulting is a Namibia-registered technology consulting firm founded and led by Job Angula, a seasoned full-stack software developer and technology consultant. The firm specialises in custom web development, application design, and digital solutions for organisations seeking purpose-built technology platforms.")
add_body(doc, "Operating from Windhoek, Namibia, the firm serves clients across the Southern African region, with particular expertise in developing solutions for organisations that require robust, secure, and user-friendly digital platforms.")

doc.add_paragraph("")

# Core Competencies - Visual cards
add_section_heading(doc, "2. Core Competencies", level=2)
doc.add_paragraph("")

add_icon_card_row(doc, [
    ("\U0001F310", "Web Development", "Custom websites &\nweb applications"),
    ("\U0001F3A8", "UI/UX Design", "User-centred design\nfor intuitive interfaces"),
    ("\U0001F4BB", "Custom CMS", "Purpose-built content\nmanagement systems"),
])

doc.add_paragraph("")

add_icon_card_row(doc, [
    ("\u2601\ufe0f", "Cloud Hosting", "Reliable infrastructure\n& server management"),
    ("\U0001F512", "Security", "Data protection &\ncyber security"),
    ("\U0001F30D", "Multilingual", "Bilingual/multilingual\nweb solutions"),
])

doc.add_paragraph("")

add_section_heading(doc, "3. Principal Consultant", level=2)
doc.add_paragraph("")

add_info_box(doc,
    "Job Angula \u2013 Principal Consultant & Lead Developer",
    "Full-stack web developer and technology consultant with extensive experience in designing and "
    "developing web-based solutions for diverse clients. Technical proficiency spans front-end design, "
    "back-end development, database architecture, server administration, and cloud hosting.\n\n"
    "Key Skills: HTML5, CSS3, JavaScript, React/Next.js, Python, Node.js, PostgreSQL, REST APIs, "
    "AWS, DigitalOcean, Web Security, WCAG 2.1 Accessibility\n\n"
    "[ATTACH FULL CV AS ANNEXURE]",
    SECTION_BG
)

doc.add_paragraph("")

add_section_heading(doc, "4. Specialist Subcontractors", level=2)
add_body(doc, "For specialised deliverables, Job Angula Technology Consulting engages trusted subcontractors:")
doc.add_paragraph("")

add_styled_table(doc,
    ["Role", "Responsibility"],
    [
        ["UI/UX Designer", "Visual design concepts, wireframes, and prototyping"],
        ["Content Strategist", "Content architecture, migration planning, and information design"],
        ["QA/Testing Specialist", "Cross-browser testing, performance testing, and quality assurance"],
        ["Portuguese Language QA", "Linguistic quality assurance for the Portuguese interface"],
    ],
    [5, 11]
)

doc.add_paragraph("")

add_section_heading(doc, "5. Business Details", level=2)
doc.add_paragraph("")

add_styled_table(doc,
    ["Detail", "Information"],
    [
        ["Business Name", "Job Angula Technology Consulting"],
        ["Principal", "Job Angula"],
        ["Location", "Windhoek, Namibia"],
        ["Phone", "+264 XX XXX XXXX"],
        ["Email", "job@angulaconsulting.com"],
        ["Website", "www.angulaconsulting.com"],
        ["Registration", "[INSERT BUSINESS REGISTRATION NUMBER]"],
    ],
    [5, 11]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
#                    SECTION C: TECHNICAL PROPOSAL
# ═══════════════════════════════════════════════════════════════════

add_section_heading(doc, "Section C: Technical Proposal")
doc.add_paragraph("")

# ── Understanding of TOR ──
add_section_heading(doc, "1. Understanding of the Terms of Reference", level=2)

add_body(doc, "The Benguela Current Convention (BCC) is a pioneering intergovernmental organisation established by Angola, Namibia, and South Africa to promote integrated management and sustainable development of the Benguela Current Large Marine Ecosystem. As the first convention in the world based on a multi-sectoral approach to Large Marine Ecosystem ocean governance, BCC occupies a unique and prestigious position in the global marine conservation landscape.")

add_body(doc, "The BCC Secretariat communicates with a diverse stakeholder base \u2013 including policy makers, administrators, technocrats, academics, students, the private sector, and the general public \u2013 in both English and Portuguese. A modern, professional website is central to this communication mandate.")

doc.add_paragraph("")

# Key Requirements Visual
add_section_heading(doc, "Core Requirements Identified:", level=3)
doc.add_paragraph("")

req_table = doc.add_table(rows=4, cols=2)
req_table.alignment = WD_TABLE_ALIGNMENT.CENTER

requirements = [
    ("\U0001F310 Website Revamp", "Modern, interactive redesign reflecting BCC\u2019s identity"),
    ("\U0001F4AC Bilingual Platform", "Full English & Portuguese support with seamless switching"),
    ("\u2699\ufe0f Custom CMS", "Intuitive system for Secretariat staff to manage content independently"),
    ("\U0001F4C1 Document Repository", "Secure, searchable repository for reports and policy documents"),
    ("\U0001F91D Stakeholder Engagement", "Social media integration, newsletter, and interactive features"),
    ("\u2601\ufe0f Hosting & Maintenance", "Reliable 3-year hosting for BCC website and BCLME RIIMS"),
    ("\U0001F393 Training & Handover", "Comprehensive staff training with user manuals"),
    ("\U0001F512 Compliance", "Data protection and WCAG 2.1 accessibility standards"),
]

for i in range(4):
    for j in range(2):
        idx = i * 2 + j
        cell = req_table.rows[i].cells[j]
        cell.text = ""
        set_cell_shading(cell, SECTION_BG.replace("#", ""))
        borders = {"top": {"val": "single", "sz": "2", "color": "CCCCCC"},
                   "bottom": {"val": "single", "sz": "2", "color": "CCCCCC"},
                   "left": {"val": "single", "sz": "2", "color": "CCCCCC"},
                   "right": {"val": "single", "sz": "2", "color": "CCCCCC"}}
        set_cell_border(cell, **borders)

        p_req_title = cell.paragraphs[0]
        run_rt = p_req_title.add_run(requirements[idx][0])
        run_rt.bold = True
        run_rt.font.size = Pt(10)
        run_rt.font.color.rgb = DEEP_BLUE
        run_rt.font.name = 'Calibri'

        p_req_desc = cell.add_paragraph()
        run_rd = p_req_desc.add_run(requirements[idx][1])
        run_rd.font.size = Pt(9)
        run_rd.font.color.rgb = MEDIUM_GRAY
        run_rd.font.name = 'Calibri'

doc.add_paragraph("")

# ── Why Custom CMS ──
add_section_heading(doc, "2. Proposed Technical Approach", level=2)
doc.add_paragraph("")

add_section_heading(doc, "2.1 Why a Custom CMS?", level=3)
add_body(doc, "Rather than relying on off-the-shelf platforms such as WordPress or Drupal, we propose building a bespoke Content Management System specifically engineered for BCC\u2019s needs. This approach offers significant advantages:")
doc.add_paragraph("")

add_styled_table(doc,
    ["Factor", "Custom CMS (Our Approach)", "Off-the-shelf (WordPress)"],
    [
        ["Security", "Minimal attack surface; no known public exploits", "Frequent target of automated attacks; thousands of known vulnerabilities"],
        ["Performance", "Lean codebase; only features BCC needs", "Bloated with unused features and plugins"],
        ["Tailored UX", "Admin interface designed for BCC staff workflows", "Generic admin panels requiring adaptation"],
        ["Maintenance", "No dependency on third-party plugin updates", "Constant plugin and core updates required"],
        ["Ownership", "BCC owns 100% of the code", "Dependent on open-source community"],
        ["Bilingual", "Native dual-language architecture from ground up", "Relies on third-party translation plugins"],
    ],
    [3, 6.5, 6.5]
)

doc.add_paragraph("")

# ── Technology Stack ──
add_section_heading(doc, "2.2 Technology Stack", level=3)
doc.add_paragraph("")

add_styled_table(doc,
    ["Component", "Technology", "Rationale"],
    [
        ["Front-end", "HTML5, CSS3, React / Next.js", "Modern, fast, responsive UI with excellent SEO"],
        ["Back-end", "Node.js / Python (Django)", "Robust, scalable server-side framework"],
        ["Database", "PostgreSQL", "Enterprise-grade; excellent for multilingual content"],
        ["CMS Admin", "Custom-built dashboard", "Intuitive interface tailored to BCC workflows"],
        ["Hosting", "Cloud VPS (AWS / DigitalOcean)", "High availability; data centres near Southern Africa"],
        ["CDN", "Cloudflare", "Global content delivery for fast page loads"],
        ["Security", "SSL + WAF + CSP Headers", "HTTPS encryption and web application firewall"],
        ["Analytics", "Built-in + Google Analytics", "Hit counter and detailed traffic analytics"],
    ],
    [3, 5, 8]
)

doc.add_paragraph("")

# ── Three Design Concepts ──
add_section_heading(doc, "2.3 Three Design Concepts", level=3)
add_body(doc, "As required, we will deliver three distinct design concepts for BCC\u2019s consideration:")
doc.add_paragraph("")

# Concept cards
concepts = [
    ("Concept 1: \u201cOcean Authority\u201d",
     "0073B7",
     "A bold, authoritative design with deep ocean blues and clean lines. Emphasises BCC\u2019s role as a "
     "serious intergovernmental body. Features large hero imagery of the Benguela Current ecosystem, "
     "structured navigation, and a formal yet accessible layout."),
    ("Concept 2: \u201cLiving Ecosystem\u201d",
     "009B72",
     "A vibrant, dynamic design inspired by the marine biodiversity of the BCLME. Uses rich photography, "
     "animated ocean elements, and a colour palette drawn from the sea. Interactive maps and visual "
     "storytelling are central features."),
    ("Concept 3: \u201cConnected Shores\u201d",
     "D4A017",
     "A warm, collaborative design emphasising the partnership between Angola, Namibia, and South Africa. "
     "Features the three nations\u2019 colours as accent tones within BCC\u2019s brand palette. Clean, minimal "
     "layout with emphasis on readability and accessibility."),
]

for title, color, desc in concepts:
    c_table = doc.add_table(rows=1, cols=1)
    c_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_cell = c_table.rows[0].cells[0]
    set_cell_shading(c_cell, "FFFFFF")
    borders = {"top": {"val": "single", "sz": "2", "color": color},
               "bottom": {"val": "single", "sz": "2", "color": color},
               "left": {"val": "single", "sz": "18", "color": color},
               "right": {"val": "single", "sz": "2", "color": color}}
    set_cell_border(c_cell, **borders)
    c_cell.text = ""
    p_ct = c_cell.paragraphs[0]
    run_ct = p_ct.add_run(title)
    run_ct.bold = True
    run_ct.font.size = Pt(12)
    run_ct.font.color.rgb = RGBColor(int(color[0:2],16), int(color[2:4],16), int(color[4:6],16))
    run_ct.font.name = 'Calibri'

    p_cd = c_cell.add_paragraph()
    run_cd = p_cd.add_run(desc)
    run_cd.font.size = Pt(10)
    run_cd.font.color.rgb = DARK_GRAY
    run_cd.font.name = 'Calibri'

    doc.add_paragraph("")  # spacing

add_body(doc, "All three concepts will be fully responsive (desktop, tablet, mobile), compliant with WCAG 2.1 accessibility standards, and reflective of BCC\u2019s corporate identity.")

doc.add_paragraph("")

# ── Key Features ──
add_section_heading(doc, "2.4 Key Features & Functionality", level=3)
doc.add_paragraph("")

features = [
    ("\U0001F4AC", "Bilingual Architecture", "EN/PT",
     "Language toggle in header; side-by-side content editing; SEO-friendly URL prefixes (/en/, /pt/)"),
    ("\u2699\ufe0f", "Custom CMS", "Admin Panel",
     "WYSIWYG editor, media library, menu manager, role-based access, content scheduling, revision history"),
    ("\U0001F4C1", "Document Repository", "Secure",
     "Categorised library with search/filter, access controls, bulk upload, download tracking"),
    ("\U0001F4F0", "News & Events", "Dynamic",
     "Article publishing, events calendar, conference/meeting registration functionality"),
    ("\U0001F4E7", "Newsletter", "Mailing List",
     "Subscription with email verification, subscriber management, SendGrid/Mailchimp integration"),
    ("\U0001F4F1", "Social Media", "Integration",
     "Live feed widgets, share buttons, Open Graph meta tags for Facebook, LinkedIn, Instagram, YouTube, X, TikTok"),
    ("\U0001F4DD", "Sub-pages & Forms", "Dynamic",
     "Template-based page creation for projects, tenders, job applications; custom forms builder"),
    ("\U0001F4CA", "Analytics", "Built-in",
     "Visitor dashboard with geographic data, exportable reports, Google Analytics integration"),
    ("\U0001F50D", "Search", "Full-text",
     "Cross-content search with filters by type, language, date; search analytics"),
]

for icon, name, badge, desc in features:
    f_table = doc.add_table(rows=1, cols=1)
    f_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    f_cell = f_table.rows[0].cells[0]
    set_cell_shading(f_cell, SECTION_BG.replace("#",""))
    borders = {"left": {"val": "single", "sz": "12", "color": ACCENT_BAR}}
    set_cell_border(f_cell, **borders)
    f_cell.text = ""

    p_fn = f_cell.paragraphs[0]
    run_fi = p_fn.add_run(icon + "  ")
    run_fi.font.size = Pt(12)
    run_fname = p_fn.add_run(name)
    run_fname.bold = True
    run_fname.font.size = Pt(11)
    run_fname.font.color.rgb = DEEP_BLUE
    run_fname.font.name = 'Calibri'
    run_badge = p_fn.add_run(f"  [{badge}]")
    run_badge.font.size = Pt(9)
    run_badge.font.color.rgb = OCEAN_BLUE
    run_badge.font.name = 'Calibri'

    p_fd = f_cell.add_paragraph()
    run_fd = p_fd.add_run(desc)
    run_fd.font.size = Pt(10)
    run_fd.font.color.rgb = DARK_GRAY
    run_fd.font.name = 'Calibri'


doc.add_paragraph("")

# ── Content Migration ──
add_section_heading(doc, "2.5 Content Migration Strategy", level=3)
doc.add_paragraph("")

migration_steps = [
    ("1", "AUDIT", "Catalogue all existing content (pages, documents, images, media)"),
    ("2", "PRIORITISE", "Work with BCC to identify content to migrate, update, or retire"),
    ("3", "MIGRATE", "Transfer all approved content with proper categorisation"),
    ("4", "VERIFY", "Cross-check migrated content for accuracy and completeness"),
    ("5", "REDIRECT", "Implement URL redirects to preserve SEO and prevent broken links"),
]

mig_table = doc.add_table(rows=1, cols=5)
mig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
colors_list = ["003D6B", "005A9C", "0073B7", "009B72", "00B386"]

for i, (num, label, desc) in enumerate(migration_steps):
    cell = mig_table.rows[0].cells[i]
    set_cell_shading(cell, colors_list[i])
    borders = {"top": {"val": "single", "sz": "2", "color": colors_list[i]},
               "bottom": {"val": "single", "sz": "2", "color": colors_list[i]},
               "left": {"val": "single", "sz": "2", "color": colors_list[i]},
               "right": {"val": "single", "sz": "2", "color": colors_list[i]}}
    set_cell_border(cell, **borders)
    cell.text = ""

    p_num = cell.paragraphs[0]
    p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_n = p_num.add_run(num)
    run_n.bold = True
    run_n.font.size = Pt(16)
    run_n.font.color.rgb = WHITE
    run_n.font.name = 'Calibri'

    p_label = cell.add_paragraph()
    p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_l = p_label.add_run(label)
    run_l.bold = True
    run_l.font.size = Pt(8)
    run_l.font.color.rgb = WHITE
    run_l.font.name = 'Calibri'

doc.add_paragraph("")

# Descriptions below
for num, label, desc in migration_steps:
    add_bullet(doc, f" {desc}", f"Step {num} \u2013 {label}:")

doc.add_paragraph("")

# ── Security ──
add_section_heading(doc, "2.6 Security Measures", level=3)
doc.add_paragraph("")

security_items = [
    "SSL/TLS encryption (HTTPS) on all pages",
    "Web Application Firewall (WAF) via Cloudflare",
    "Regular automated security scans",
    "Input validation and sanitisation against SQL injection and XSS",
    "Secure authentication with password hashing (bcrypt)",
    "Role-based access control for CMS users",
    "Automated daily backups with off-site storage",
    "DDoS protection via Cloudflare",
    "Content Security Policy (CSP) headers",
    "Regular security patches and updates",
]

for item in security_items:
    add_bullet(doc, item, "\u2713 ")

doc.add_paragraph("")

# ── Hosting Architecture ──
add_section_heading(doc, "2.7 Hosting Architecture", level=3)
add_body(doc, "The BCC website and BCLME RIIMS will be hosted on a high-availability cloud infrastructure:")
doc.add_paragraph("")

add_styled_table(doc,
    ["Component", "Specification"],
    [
        ["Primary Server", "Cloud VPS \u2013 4+ vCPUs, 8GB+ RAM, SSD storage"],
        ["Database", "Managed PostgreSQL with automated backups"],
        ["CDN", "Cloudflare for global delivery and DDoS protection"],
        ["Backups", "Daily automated, 30-day retention, geographically separate"],
        ["Uptime SLA", "99.9% uptime guarantee"],
        ["Monitoring", "24/7 server monitoring with automated alerts"],
    ],
    [5, 11]
)

doc.add_paragraph("")

# ── Training ──
add_section_heading(doc, "2.8 Training & Handover", level=3)
doc.add_paragraph("")

add_styled_table(doc,
    ["Training Component", "Description", "Duration"],
    [
        ["CMS Administration", "Creating/editing pages, managing documents, publishing news", "1 day"],
        ["Media Management", "Uploading and organising images, videos, and documents", "0.5 day"],
        ["User & Access Mgmt", "Managing user accounts and roles", "0.5 day"],
        ["Analytics & Reporting", "Using the analytics dashboard and generating reports", "0.5 day"],
        ["Troubleshooting", "Common issues and how to resolve them", "0.5 day"],
    ],
    [4, 9, 3]
)

doc.add_paragraph("")
add_body(doc, "Deliverables: Comprehensive User Manual (printed & digital PDF), video tutorials for key CMS functions, quick-reference guide, and ongoing email/phone support during the 3-year hosting period.")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
#                 SECTION C (cont): METHODOLOGY
# ═══════════════════════════════════════════════════════════════════

add_section_heading(doc, "3. Implementation Methodology", level=2)
add_body(doc, "Our implementation follows an Agile-inspired methodology with structured phases and regular client feedback:")
doc.add_paragraph("")

phases = [
    ("1", "Discovery & Planning", "Weeks 1\u20132", [
        "Inception meeting with BCC Secretariat",
        "Audit of current website structure, content, and functionality",
        "Stakeholder requirements gathering",
        "Content inventory and migration planning",
        "Technical architecture design",
        "Inception Report delivery",
    ]),
    ("2", "Design", "Weeks 3\u20135", [
        "Development of three design concepts",
        "Presentation to BCC for review and selection",
        "Refinement of selected design",
        "Responsive layouts for desktop, tablet, mobile",
    ]),
    ("3", "Development", "Weeks 5\u20139", [
        "Custom CMS core development",
        "Front-end development (all pages and templates)",
        "Bilingual architecture implementation",
        "Document repository, newsletter, social media, forms, analytics modules",
    ]),
    ("4", "Migration & Testing", "Weeks 9\u201311", [
        "Content migration from existing website",
        "Functionality, speed, responsiveness, and security testing",
        "User acceptance testing with BCC staff",
        "Bug fixes and refinements",
    ]),
    ("5", "Launch & Training", "Weeks 11\u201312", [
        "Final deployment to production servers",
        "DNS migration, SSL setup, URL redirects",
        "Staff training sessions (3 days in Swakopmund)",
        "User manual and technical documentation delivery",
    ]),
    ("6", "Hosting & Maintenance", "Years 1\u20133", [
        "24/7 server monitoring and security patches",
        "Monthly maintenance and quarterly performance reports",
        "Visitor analytics reports on request",
        "Ongoing technical support via email and phone",
    ]),
]

phase_colors = ["003D6B", "005A9C", "0073B7", "009B72", "00B386", "D4A017"]

for i, (num, title, duration, items) in enumerate(phases):
    add_phase_block(doc, num, title, duration, items, phase_colors[i])
    doc.add_paragraph("")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
#                    SECTION D: WORK PLAN
# ═══════════════════════════════════════════════════════════════════

add_section_heading(doc, "Section D: Work Plan")
doc.add_paragraph("")

add_body(doc, "The following detailed work plan outlines tasks and deliverables across the 12-week development phase and the subsequent 3-year hosting period.")
doc.add_paragraph("")

workplan = [
    ["Week 1", "Inception meeting; stakeholder consultation; website audit", "Inception Report"],
    ["Week 2", "Requirements finalisation; content inventory; architecture planning", "Requirements Document"],
    ["Week 3", "Wireframe development; design concept 1", "Wireframes"],
    ["Week 4", "Design concepts 2 and 3", "Three Design Concepts"],
    ["Week 5", "Design presentation; feedback; refinement and approval", "Approved Final Design"],
    ["Week 6", "CMS core development; database design; bilingual setup", "Progress Report"],
    ["Week 7", "Front-end development; page templates; responsive layouts", "Progress Report"],
    ["Week 8", "Document repository; newsletter; social media integration", "Progress Report"],
    ["Week 9", "Forms, registration, analytics modules; CMS admin completion", "Progress Report"],
    ["Week 10", "Content migration; internal testing and QA", "Migration & Test Report"],
    ["Week 11", "User acceptance testing; bug fixes and refinements", "UAT Sign-off"],
    ["Week 12", "Deployment; DNS migration; staff training; documentation", "Live Website; Manuals"],
    ["Months 4\u201336", "Hosting, maintenance, security updates, and support", "Monthly/Quarterly Reports"],
]

add_styled_table(doc,
    ["Timeline", "Activity", "Deliverable"],
    workplan,
    [3, 8.5, 4.5]
)

doc.add_paragraph("")

# Visual Gantt-like chart
add_section_heading(doc, "Visual Project Timeline", level=2)
doc.add_paragraph("")

gantt_table = doc.add_table(rows=7, cols=13)
gantt_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row (week numbers)
gantt_headers = ["Phase", "W1", "W2", "W3", "W4", "W5", "W6", "W7", "W8", "W9", "W10", "W11", "W12"]
for i, h in enumerate(gantt_headers):
    cell = gantt_table.rows[0].cells[i]
    cell.text = ""
    set_cell_shading(cell, TABLE_HEADER_BG)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(7)
    r.font.color.rgb = WHITE
    r.font.name = 'Calibri'

# Phase data: (name, start_week, end_week, color)
gantt_phases = [
    ("Discovery & Planning", 1, 2, "003D6B"),
    ("Design", 3, 5, "005A9C"),
    ("Development", 5, 9, "0073B7"),
    ("Migration & Testing", 9, 11, "009B72"),
    ("Launch & Training", 11, 12, "00B386"),
    ("Hosting & Maintenance", 1, 12, "D4A017"),
]

for r_idx, (phase_name, start, end, color) in enumerate(gantt_phases):
    row = gantt_table.rows[r_idx + 1]
    # Phase name
    name_cell = row.cells[0]
    name_cell.text = ""
    p = name_cell.paragraphs[0]
    r = p.add_run(phase_name)
    r.bold = True
    r.font.size = Pt(7)
    r.font.color.rgb = DARK_GRAY
    r.font.name = 'Calibri'

    # Week cells
    for w in range(1, 13):
        cell = row.cells[w]
        cell.text = ""
        if start <= w <= end:
            if r_idx < 5:  # Not the hosting row
                set_cell_shading(cell, color)
            else:
                set_cell_shading(cell, "FFF3CD")  # Light yellow for ongoing hosting
        else:
            set_cell_shading(cell, "F8F8F8")

doc.add_paragraph("")
add_body(doc, "Reporting: Bi-weekly progress reports submitted to the Manager: Marketing & Communication and Manager: Data and Information Management, with oversight from the Executive Secretary.")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
#                    SECTION E: FINANCIAL PROPOSAL
# ═══════════════════════════════════════════════════════════════════

add_section_heading(doc, "Section E: Financial Proposal")
doc.add_paragraph("")

add_body(doc, "The financial proposal separates professional fees from other costs related to the assignment, as stipulated in the Terms of Reference. All amounts are in Namibian Dollars (NAD).")
doc.add_paragraph("")

# Professional Fees
add_section_heading(doc, "1. Professional Fees", level=2)
doc.add_paragraph("")

prof_fees = [
    ["", "Discovery & Planning", ""],
    ["1.1", "Inception meeting and stakeholder consultation", "5,000.00"],
    ["1.2", "Current website audit and requirements analysis", "8,000.00"],
    ["1.3", "Content inventory and migration planning", "5,000.00"],
    ["", "Subtotal: Discovery & Planning", "18,000.00"],
    ["", "Website Design", ""],
    ["2.1", "Three (3) design concepts (wireframes + mockups)", "20,000.00"],
    ["2.2", "Design refinement and responsive layouts", "8,000.00"],
    ["", "Subtotal: Design", "28,000.00"],
    ["", "Website Development", ""],
    ["3.1", "Custom CMS development (admin panel, editors, media library)", "35,000.00"],
    ["3.2", "Front-end development (pages, templates, responsive)", "20,000.00"],
    ["3.3", "Bilingual architecture (English/Portuguese)", "8,000.00"],
    ["3.4", "Document repository module", "8,000.00"],
    ["3.5", "Newsletter/mailing list module", "5,000.00"],
    ["3.6", "Social media integration", "3,000.00"],
    ["3.7", "Forms, registration, and sub-page creation module", "7,000.00"],
    ["3.8", "Visitor analytics module", "5,000.00"],
    ["3.9", "Search functionality", "3,000.00"],
    ["", "Subtotal: Development", "94,000.00"],
    ["", "Content Migration & Testing", ""],
    ["4.1", "Content migration from existing website", "8,000.00"],
    ["4.2", "Quality assurance, testing, and bug fixing", "7,000.00"],
    ["4.3", "Security testing and hardening", "3,000.00"],
    ["", "Subtotal: Migration & Testing", "18,000.00"],
    ["", "Training & Handover", ""],
    ["5.1", "Staff training (3 days on-site in Swakopmund)", "8,000.00"],
    ["5.2", "User manual and video tutorials", "4,000.00"],
    ["5.3", "Final technical documentation", "3,000.00"],
    ["", "Subtotal: Training & Handover", "15,000.00"],
]

pf_table = doc.add_table(rows=1 + len(prof_fees), cols=3)
pf_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
for i, h in enumerate(["No.", "Description", "Amount (NAD)"]):
    cell = pf_table.rows[0].cells[i]
    cell.text = ""
    set_cell_shading(cell, TABLE_HEADER_BG)
    p = cell.paragraphs[0]
    if i == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = WHITE
    r.font.name = 'Calibri'

for r_idx, (no, desc, amt) in enumerate(prof_fees):
    row = pf_table.rows[r_idx + 1]
    is_subtotal = desc.startswith("Subtotal")
    is_header = no == "" and not is_subtotal

    for c_idx, val in enumerate([no, desc, amt]):
        cell = row.cells[c_idx]
        cell.text = ""
        p = cell.paragraphs[0]
        if c_idx == 2:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(val)
        r.font.size = Pt(10)
        r.font.name = 'Calibri'

        if is_header:
            r.bold = True
            r.font.color.rgb = DEEP_BLUE
            set_cell_shading(cell, "E8F0FE")
        elif is_subtotal:
            r.bold = True
            r.font.color.rgb = DEEP_BLUE
            set_cell_shading(cell, "D6EAF8")
        else:
            r.font.color.rgb = DARK_GRAY
            if r_idx % 2 == 0:
                set_cell_shading(cell, TABLE_ALT_ROW)

# Total row
total_row = pf_table.add_row()
for i, val in enumerate(["", "TOTAL PROFESSIONAL FEES", "173,000.00"]):
    cell = total_row.cells[i]
    cell.text = ""
    set_cell_shading(cell, "003D6B")
    p = cell.paragraphs[0]
    if i == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(val)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = WHITE
    r.font.name = 'Calibri'

doc.add_paragraph("")

# Other Costs
add_section_heading(doc, "2. Other Costs", level=2)
doc.add_paragraph("")

other_costs = [
    ["6.1", "Cloud hosting \u2013 BCC Website (Year 1)", "12,000.00"],
    ["6.2", "Cloud hosting \u2013 BCC Website (Year 2)", "12,000.00"],
    ["6.3", "Cloud hosting \u2013 BCC Website (Year 3)", "12,000.00"],
    ["6.4", "Cloud hosting \u2013 BCLME RIIMS (Year 1)", "8,000.00"],
    ["6.5", "Cloud hosting \u2013 BCLME RIIMS (Year 2)", "8,000.00"],
    ["6.6", "Cloud hosting \u2013 BCLME RIIMS (Year 3)", "8,000.00"],
    ["6.7", "Domain and SSL certificates (3 years)", "3,000.00"],
    ["6.8", "CDN service \u2013 Cloudflare (3 years)", "0.00"],
    ["6.9", "Monthly maintenance & security (36 months x NAD 500)", "18,000.00"],
    ["7.1", "Travel to Swakopmund (2 trips)", "6,000.00"],
    ["7.2", "Accommodation (6 nights)", "6,000.00"],
]

add_styled_table(doc,
    ["No.", "Description", "Amount (NAD)"],
    other_costs,
    [2, 10, 4]
)

# Other costs total
doc.add_paragraph("")
oc_total_table = doc.add_table(rows=1, cols=3)
oc_total_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, val in enumerate(["", "TOTAL OTHER COSTS", "93,000.00"]):
    cell = oc_total_table.rows[0].cells[i]
    cell.text = ""
    set_cell_shading(cell, "003D6B")
    p = cell.paragraphs[0]
    if i == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(val)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = WHITE
    r.font.name = 'Calibri'

doc.add_paragraph("")
doc.add_paragraph("")

# Grand Summary
add_section_heading(doc, "3. Summary", level=2)
doc.add_paragraph("")

# Budget Breakdown Visual (pie chart description)
add_info_box(doc,
    "Budget Allocation Overview",
    "Professional Fees: NAD 173,000 (65%) \u2502 Hosting & Maintenance: NAD 81,000 (30%) \u2502 Travel: NAD 12,000 (5%)",
    SECTION_BG
)

doc.add_paragraph("")

# Visual budget bars
add_progress_bar(doc, "Professional Fees", 65, ACCENT_BAR)
add_progress_bar(doc, "Hosting (3 Yrs)", 30, "009B72")
add_progress_bar(doc, "Travel", 5, "D4A017")

doc.add_paragraph("")

summary_table = doc.add_table(rows=4, cols=2)
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

summary_data = [
    ("Professional Fees", "173,000.00"),
    ("Other Costs (Hosting, Maintenance, Travel)", "93,000.00"),
    ("TOTAL (VAT Exclusive)", "266,000.00"),
    ("VAT (15%)", "39,900.00"),
]

for i, (label, amount) in enumerate(summary_data):
    cell0 = summary_table.rows[i].cells[0]
    cell1 = summary_table.rows[i].cells[1]

    cell0.text = ""
    cell1.text = ""

    p0 = cell0.paragraphs[0]
    p1 = cell1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    r0 = p0.add_run(label)
    r1 = p1.add_run("NAD " + amount)

    r0.font.size = Pt(11)
    r1.font.size = Pt(11)
    r0.font.name = 'Calibri'
    r1.font.name = 'Calibri'

    if i >= 2:
        r0.bold = True
        r1.bold = True
        r0.font.color.rgb = DEEP_BLUE
        r1.font.color.rgb = DEEP_BLUE
        set_cell_shading(cell0, "D6EAF8")
        set_cell_shading(cell1, "D6EAF8")
    else:
        r0.font.color.rgb = DARK_GRAY
        r1.font.color.rgb = DARK_GRAY
        if i % 2 == 0:
            set_cell_shading(cell0, TABLE_ALT_ROW)
            set_cell_shading(cell1, TABLE_ALT_ROW)

# Grand total
gt_row = summary_table.add_row()
gt_cell0 = gt_row.cells[0]
gt_cell1 = gt_row.cells[1]
gt_cell0.text = ""
gt_cell1.text = ""
set_cell_shading(gt_cell0, "003D6B")
set_cell_shading(gt_cell1, "003D6B")

p_gt0 = gt_cell0.paragraphs[0]
r_gt0 = p_gt0.add_run("GRAND TOTAL (VAT Inclusive)")
r_gt0.bold = True
r_gt0.font.size = Pt(12)
r_gt0.font.color.rgb = WHITE
r_gt0.font.name = 'Calibri'

p_gt1 = gt_cell1.paragraphs[0]
p_gt1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r_gt1 = p_gt1.add_run("NAD 305,900.00")
r_gt1.bold = True
r_gt1.font.size = Pt(12)
r_gt1.font.color.rgb = WHITE
r_gt1.font.name = 'Calibri'

doc.add_paragraph("")

add_info_box(doc,
    "Note",
    "If Job Angula Technology Consulting is not VAT-registered, please remove the VAT line. "
    "The Grand Total would then be NAD 266,000.00. Adjust accordingly before submission.",
    "FFF8E1"
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
#                    SECTION F: REFERENCES
# ═══════════════════════════════════════════════════════════════════

add_section_heading(doc, "Section F: References")
doc.add_paragraph("")

# Reference 1
ref1_table = doc.add_table(rows=1, cols=1)
ref1_table.alignment = WD_TABLE_ALIGNMENT.CENTER
ref1_cell = ref1_table.rows[0].cells[0]
set_cell_shading(ref1_cell, SECTION_BG.replace("#",""))
borders = {"left": {"val": "single", "sz": "18", "color": "003D6B"}}
set_cell_border(ref1_cell, **borders)
ref1_cell.text = ""

p_r1t = ref1_cell.paragraphs[0]
run_r1t = p_r1t.add_run("Reference 1: Angula Consulting Website")
run_r1t.bold = True
run_r1t.font.size = Pt(12)
run_r1t.font.color.rgb = DEEP_BLUE
run_r1t.font.name = 'Calibri'

ref1_details = [
    ("Project:", "Design and development of the Angula Consulting corporate website"),
    ("URL:", "www.angulaconsulting.com"),
    ("Description:", "Full design and development of a professional consulting firm website showcasing services, portfolio, and contact capabilities. Built with modern responsive design principles and optimised for performance and SEO."),
    ("Year:", "[INSERT YEAR]"),
    ("Technologies:", "[INSERT: e.g., React, HTML/CSS/JS, etc.]"),
]

for label, value in ref1_details:
    p = ref1_cell.add_paragraph()
    rl = p.add_run(label + " ")
    rl.bold = True
    rl.font.size = Pt(10)
    rl.font.color.rgb = OCEAN_BLUE
    rl.font.name = 'Calibri'
    rv = p.add_run(value)
    rv.font.size = Pt(10)
    rv.font.color.rgb = DARK_GRAY
    rv.font.name = 'Calibri'

doc.add_paragraph("")

# Reference 2
ref2_table = doc.add_table(rows=1, cols=1)
ref2_table.alignment = WD_TABLE_ALIGNMENT.CENTER
ref2_cell = ref2_table.rows[0].cells[0]
set_cell_shading(ref2_cell, SECTION_BG.replace("#",""))
borders2 = {"left": {"val": "single", "sz": "18", "color": "0073B7"}}
set_cell_border(ref2_cell, **borders2)
ref2_cell.text = ""

p_r2t = ref2_cell.paragraphs[0]
run_r2t = p_r2t.add_run("Reference 2: OndyPOS Website & System")
run_r2t.bold = True
run_r2t.font.size = Pt(12)
run_r2t.font.color.rgb = DEEP_BLUE
run_r2t.font.name = 'Calibri'

ref2_details = [
    ("Project:", "Design and development of the OndyPOS point-of-sale system and website"),
    ("URL:", "www.ondypos.com"),
    ("Description:", "Full-stack development of OndyPOS, a comprehensive point-of-sale and inventory management web application, along with its marketing website. Demonstrates expertise in custom application development, database design, user authentication, and intuitive admin interfaces \u2013 skills directly transferable to the BCC custom CMS requirement."),
    ("Year:", "[INSERT YEAR]"),
    ("Contact:", "[INSERT CLIENT CONTACT NAME, PHONE, EMAIL]"),
    ("Technologies:", "[INSERT: e.g., React, Node.js, PostgreSQL, etc.]"),
]

for label, value in ref2_details:
    p = ref2_cell.add_paragraph()
    rl = p.add_run(label + " ")
    rl.bold = True
    rl.font.size = Pt(10)
    rl.font.color.rgb = OCEAN_BLUE
    rl.font.name = 'Calibri'
    rv = p.add_run(value)
    rv.font.size = Pt(10)
    rv.font.color.rgb = DARK_GRAY
    rv.font.name = 'Calibri'

doc.add_paragraph("")

# Reference 3
ref3_table = doc.add_table(rows=1, cols=1)
ref3_table.alignment = WD_TABLE_ALIGNMENT.CENTER
ref3_cell = ref3_table.rows[0].cells[0]
set_cell_shading(ref3_cell, SECTION_BG.replace("#",""))
borders3 = {"left": {"val": "single", "sz": "18", "color": "009B72"}}
set_cell_border(ref3_cell, **borders3)
ref3_cell.text = ""

p_r3t = ref3_cell.paragraphs[0]
run_r3t = p_r3t.add_run("Reference 3: GovRecruit Namibia Website")
run_r3t.bold = True
run_r3t.font.size = Pt(12)
run_r3t.font.color.rgb = DEEP_BLUE
run_r3t.font.name = 'Calibri'

ref3_details = [
    ("Project:", "Design and development of the GovRecruit Namibia platform"),
    ("URL:", "[INSERT URL]"),
    ("Description:", "Design and development of the GovRecruit Namibia platform, a web-based government recruitment and job application system. Involved building user-facing pages, application submission forms, document upload functionality, and an administrative back-end \u2013 directly relevant to BCC\u2019s requirement for sub-pages for job applications, tender submissions, and conference registrations."),
    ("Year:", "[INSERT YEAR]"),
    ("Contact:", "[INSERT CLIENT CONTACT NAME, PHONE, EMAIL]"),
    ("Technologies:", "[INSERT: e.g., React, Django, PostgreSQL, etc.]"),
]

for label, value in ref3_details:
    p = ref3_cell.add_paragraph()
    rl = p.add_run(label + " ")
    rl.bold = True
    rl.font.size = Pt(10)
    rl.font.color.rgb = OCEAN_BLUE
    rl.font.name = 'Calibri'
    rv = p.add_run(value)
    rv.font.size = Pt(10)
    rv.font.color.rgb = DARK_GRAY
    rv.font.name = 'Calibri'

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
#                    ANNEXURES CHECKLIST
# ═══════════════════════════════════════════════════════════════════

add_section_heading(doc, "Annexures: Supporting Documents")
doc.add_paragraph("")

add_body(doc, "The following supporting documents are to be attached to this proposal:")
doc.add_paragraph("")

annexures = [
    ("Annexure 1", "Curriculum Vitae of Job Angula (Principal Consultant)", "\u2610"),
    ("Annexure 2", "Certified copy of qualifications", "\u2610"),
    ("Annexure 3", "Business registration certificate", "\u2610"),
    ("Annexure 4", "Reference letters / testimonials from clients", "\u2610"),
    ("Annexure 5", "Portfolio screenshots of referenced projects", "\u2610"),
]

ann_table = doc.add_table(rows=len(annexures), cols=3)
ann_table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, (annex, desc, check) in enumerate(annexures):
    cell0 = ann_table.rows[i].cells[0]
    cell1 = ann_table.rows[i].cells[1]
    cell2 = ann_table.rows[i].cells[2]

    cell0.text = ""
    cell1.text = ""
    cell2.text = ""

    p0 = cell0.paragraphs[0]
    r0 = p0.add_run(annex)
    r0.bold = True
    r0.font.size = Pt(10)
    r0.font.color.rgb = OCEAN_BLUE
    r0.font.name = 'Calibri'

    p1 = cell1.paragraphs[0]
    r1 = p1.add_run(desc)
    r1.font.size = Pt(10)
    r1.font.color.rgb = DARK_GRAY
    r1.font.name = 'Calibri'

    p2 = cell2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(check)
    r2.font.size = Pt(14)
    r2.font.color.rgb = MEDIUM_GRAY

    if i % 2 == 0:
        set_cell_shading(cell0, TABLE_ALT_ROW)
        set_cell_shading(cell1, TABLE_ALT_ROW)
        set_cell_shading(cell2, TABLE_ALT_ROW)

doc.add_paragraph("")
doc.add_paragraph("")

# Final footer
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr_f = p_footer._p.get_or_add_pPr()
pBdr_f = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:top w:val="single" w:sz="12" w:space="8" w:color="{ACCENT_BAR}"/>'
    f'</w:pBdr>'
)
pPr_f.append(pBdr_f)

run_footer = p_footer.add_run(
    "This proposal is submitted by Job Angula Technology Consulting in response to the "
    "BCC Call for Expression of Interest for the Revamping and Re-Designing of the BCC Website, "
    "February 2026."
)
run_footer.italic = True
run_footer.font.size = Pt(9)
run_footer.font.color.rgb = MEDIUM_GRAY
run_footer.font.name = 'Calibri'


# ═══════════════════════════════════════════════════════════════════
#                    SAVE
# ═══════════════════════════════════════════════════════════════════

output_path = "/Users/angula/Desktop/benguela/BCC_Website_Proposal_Job_Angula_Technology_Consulting.docx"
doc.save(output_path)
print(f"Proposal saved to: {output_path}")
print("Done!")
